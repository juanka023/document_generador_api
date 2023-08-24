import io
import zipfile
from docx import Document
from django.http import HttpResponse
from django.template import Context, Template
from rest_framework.views import APIView


class DocumentGenerationView(APIView):

    def generate_text_from_template(self, context):
        txt_template = "Se remite a Sr(a) {{nombre}} la disposición de presentarse en la entidad {{entidad}} con su vehículo de placa {{placa}}."          
        txt_template = Template(txt_template)
        generated_text = txt_template.render(Context(context))

        return generated_text


    def post(self, request):
        data = request.data
        input_file = data.get('input_file', None)
        context = {}
        generated_text = ""

        if input_file:
            if input_file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                docx_template = Document(io.BytesIO(input_file.read()))
                cont_text = ""

                for paragraph in docx_template.paragraphs:
                    cont_text += paragraph.text + "\n"
                    
                for line in cont_text.splitlines():
                    parts = line.split(':', 1)
                    if len(parts) == 2:
                        key, value = parts
                        context[key.strip()] = value.strip()
                
                generated_text = self.generate_text_from_template(context)
            else:
                content = input_file.read().decode('utf-8')
                for line in content.splitlines():
                    key, value = line.split(':', 1)
                    context[key.strip()] = value.strip()
                generated_text = self.generate_text_from_template(context)
        else:
            context = {
                "nombre": data.get('nombre', ''),
                "placa": data.get('placa', ''),
                "entidad": data.get('entidad', ''),
            }

            generated_text = self.generate_text_from_template(context)

        txt_file = io.BytesIO()
        txt_file.write(generated_text.encode('utf-8'))
        txt_file.seek(0)

        docx_template = Document()
        docx_template.add_paragraph(generated_text)
        docx_file = io.BytesIO()
        docx_template.save(docx_file)
        docx_file.seek(0)

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            zip_file.writestr('GeneratedFile.txt', txt_file.getvalue())
            zip_file.writestr('GeneratedFile.docx', docx_file.getvalue())

        response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
        response['Content-Disposition'] = 'attachment;filename=GeneratedFiles.zip'

        return response
